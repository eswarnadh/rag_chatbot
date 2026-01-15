from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_core.messages import (
    BaseMessage,
    HumanMessage,
    AIMessage,
    SystemMessage
)
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredExcelLoader,
)
from pathlib import Path
import os
from pydantic import SecretStr
from langchain_community.embeddings import HuggingFaceEmbeddings
from config import settings
from docx import Document as DocxDocument

DOCS_PATH = Path(__file__).parent / "docs"
CHROMA_DB_PATH = Path(__file__).parent / "chroma_db"
print(DOCS_PATH)

chroma_txt,chroma_excel = None,None

def load_docx(file_path):
    """
    Custom loader for .docx files using python-docx.
    This is more efficient than using external packages.
    """
    
    # Open the Word document
    doc = DocxDocument(file_path)
    
    # Extract all text from paragraphs
    text = "\n".join([para.text for para in doc.paragraphs])
    
    # Return as LangChain Document
    return [Document(page_content=text, metadata={"source": str(file_path)})]

def load_documents_by_type():
    """Load documents based on file type"""
    txt_documents = []
    excel_documents= []
    
    # Loop through docs folder
    for file_path in DOCS_PATH.glob("*"):
        try:
            if file_path.suffix.lower() == ".txt":
                # Use TextLoader for .txt
                loader = TextLoader(str(file_path))
                docs = loader.load()
                txt_documents.extend(docs)
                print(f"Loaded {len(docs)} documents from {file_path.name}")
            
            elif file_path.suffix.lower() == ".docx":
                # Use custom loader for .docx
                docs = load_docx(str(file_path))
                txt_documents.extend(docs)
                print(f"Loaded {len(docs)} documents from {file_path.name}")
            
            elif file_path.suffix.lower() == ".pdf":
                # Use PyPDFLoader for .pdf
                loader = PyPDFLoader(str(file_path))
                docs = loader.load()
                txt_documents.extend(docs)
                print(f"Loaded {len(docs)} documents from {file_path.name}")

            elif file_path.suffix.lower() in ['.xls', '.xlsx']:
                # Use UnstructuredExcelLoader for Excel files
                loader = UnstructuredExcelLoader(str(file_path))
                docs = loader.load()
                excel_documents.extend(docs)
                print(f"Loaded {len(docs)} documents from {file_path.name}")
            
            else:
                print(f"Skipping unsupported file type: {file_path.name}")
            
        except Exception as e:
            print(f"Error loading {file_path.name}: {e}")
            continue
    print(len(txt_documents), len(excel_documents))
    
    return txt_documents, excel_documents

def process_new_document(file_path : str):
    global chroma_txt , chroma_excel

    file_path_obj = Path(file_path)
    print(f"Processing new document: {file_path_obj.name}")

    try:
        if file_path_obj.suffix.lower() == '.txt':
            loader = TextLoader(str(file_path))
            docs = loader.load()
            doc_type = "text"
            print(f"Loaded {len(docs)} text documents")

        elif file_path_obj.suffix.lower() == '.docx':
            docs= load_docx(str(file_path))
            doc_type= "text"
            print(f"Loaded {len(docs)} text documents")

        elif file_path_obj.suffix.lower() in ['.xlx','.xlsx']:
            loader = UnstructuredExcelLoader(str(file_path))
            docs = loader.load()
            doc_type = "excel"
            print(f"Loaded {len(docs)} text documents")

        elif file_path_obj.suffix.lower() == '.pdf':
            loader= PyPDFLoader(str(file_path))
            docs= loader.load()
            doc_type= "text"
            print(f"Loaded {len(docs)} text documents")

        else:
            raise ValueError(f"Unsupported File Type")


    except Exception as e:
        print(f"Error loading the document: {str(e)}")
        raise ValueError(f"Error loading the document: {str(e)}")
    

    if doc_type == "text":
        chunks = split_into_chunks(docs)
        print(f"created {len(chunks)} chunks")

        if chroma_txt is None:
            print("initializing text vector store")
            embeddings = HuggingFaceEmbeddings(model_name = "all-mpnet-base-v2")
            chroma_txt = Chroma.from_documents(
                documents=chunks,
                embedding=embeddings,
                persist_directory=str(CHROMA_DB_PATH/"text")
            )
        else:
            print("Adding into existing text vector store")
            chroma_txt.add_documents(chunks)
        return {
            "status" :"success",
            "chunks": len(chunks),
            "type": doc_type,
            "filename":file_path_obj.name
            }

    else:
        # Excel documents - no chunking
        print(f"Processing {len(docs)} excel documents")
        
        if chroma_excel is None:
            print("Initializing excel vector store...")
            embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
            chroma_excel = Chroma.from_documents(
                documents=docs,
                embedding=embeddings,
                persist_directory=str(CHROMA_DB_PATH/"excel")
            )
        else:
            print("Adding to existing excel vector store...")
            chroma_excel.add_documents(docs)
            
        return {
            "status": "success",
            "chunks": len(docs),
            "type": doc_type,
            "filename": file_path_obj.name
        }

        

def split_into_chunks(documents):
    splitter=RecursiveCharacterTextSplitter(
        chunk_size = 500,
        chunk_overlap = 50,
        separators=["\n\n"," ","\n",""]
    )
    chunks = splitter.split_documents(documents)
    print(f"Created {len(chunks)} chunks from documents")
    return chunks

def initialize_chroma():
    global chroma_excel
    global chroma_txt

    print("Loading documents...")
    txt_documents, excel_documents = load_documents_by_type()  
    print(f"Documents loaded: {len(txt_documents), len(excel_documents)}") 
    
    if not txt_documents and not excel_documents :
        print("no docs found")
        return None
    
    print("Splitting into chunks...")  
    chunks = split_into_chunks(txt_documents)
    print(f"Chunks created: {len(chunks)}")  
    
    print("Creating embeddings...")  
    txt_embeddings = HuggingFaceEmbeddings(model_name="all-mpnet-base-v2")
    excel_embeddings = HuggingFaceEmbeddings(model_name ="all-MiniLM-L6-v2")

    
    print("Initializing Chroma...")  
    chroma_txt = Chroma.from_documents(
        documents=chunks,
        embedding=txt_embeddings,
        persist_directory=str(CHROMA_DB_PATH/"text")
    )
    chroma_excel = Chroma.from_documents(
        documents=excel_documents,
        embedding=excel_embeddings,
        persist_directory=str(CHROMA_DB_PATH/"excel")
    )
    print("vector DB initilized successfully")
    return chroma_txt, chroma_excel

def search_similar_documents(query: str, k: int = 3):
    results = []

    if chroma_txt:
        results += chroma_txt.similarity_search(query, k=3)

    if chroma_excel:
        results += chroma_excel.similarity_search(query, k=2)

    return results


def get_context_for_query(query: str, k: int = 3):
    """
    Get formatted context from documents to send with ChatGPT.
    
    Returns a string like:
    "Based on the documents:
    
    [Document 1 content]
    
    [Document 2 content]
"
    """
    results = search_similar_documents(query, k=k)
    
    if not results:
        return "No relevant documents found."
    
    # Format results into a string
    context = "Based on the following documents:\n\n"
    
    for i, doc in enumerate(results, 1):
        context += f"Document {i}:\n{doc.page_content}\n\n"
    
    return context

txt_documents, excel_documents = load_documents_by_type() 
split_into_chunks(txt_documents)